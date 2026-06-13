from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import os

# Paths
workspace = os.path.dirname(__file__)
img_path = os.path.join(workspace, 'assets', 'img', 'infografis_word2019.png')
docx_path = os.path.join(workspace, 'assets', 'files', 'Infografis_Word2019_OnePage.docx')
photo_path = os.path.join(workspace, 'assets', 'img', 'instruktur.jpg')

# Create infographic image
width, height = 1200, 1600
bg_color = (245, 250, 255)
image = Image.new('RGB', (width, height), bg_color)
draw = ImageDraw.Draw(image)

# Add background shapes
for i, rgba in enumerate([(230, 240, 255, 180), (210, 230, 250, 180), (240, 245, 255, 180)]):
    draw.rectangle([40 + i * 40, 40 + i * 40, width - 40 - i * 40, 420 + i * 40], outline=None, fill=rgba)

# Title WordArt-like effect
try:
    title_font = ImageFont.truetype('arialbd.ttf', 72)
    subtitle_font = ImageFont.truetype('arial.ttf', 28)
except IOError:
    title_font = ImageFont.load_default()
    subtitle_font = ImageFont.load_default()
title = 'Infografis Struktur Proyek'
shadow_offset = 4
x_title, y_title = 80, 80
for offset in [(shadow_offset, shadow_offset), (-shadow_offset, shadow_offset), (shadow_offset, -shadow_offset), (-shadow_offset, -shadow_offset)]:
    draw.text((x_title + offset[0], y_title + offset[1]), title, font=title_font, fill=(120, 140, 180))
draw.text((x_title, y_title), title, font=title_font, fill=(10, 40, 90))

# Subtitle
subtitle = 'SmartArt Hierarchy 3 Level + Shapes + Gambar + Text Box'
draw.text((x_title, y_title + 90), subtitle, font=subtitle_font, fill=(40, 70, 110))

# SmartArt hierarchy shapes
levels = [
    {'text': 'Level 1\nChief Project', 'xy': (420, 240, 780, 340), 'fill': (255, 255, 210)},
    {'text': 'Level 2\nTeam A', 'xy': (180, 420, 460, 520), 'fill': (201, 231, 255)},
    {'text': 'Level 2\nTeam B', 'xy': (520, 420, 800, 520), 'fill': (201, 231, 255)},
    {'text': 'Level 2\nTeam C', 'xy': (860, 420, 1140, 520), 'fill': (201, 231, 255)},
    {'text': 'Level 3\nDesigner', 'xy': (240, 620, 380, 700), 'fill': (235, 255, 235)},
    {'text': 'Level 3\nDeveloper', 'xy': (520, 620, 660, 700), 'fill': (235, 255, 235)},
    {'text': 'Level 3\nTester', 'xy': (860, 620, 1000, 700), 'fill': (235, 255, 235)},
]
for shape in levels:
    draw.rounded_rectangle(shape['xy'], radius=20, fill=shape['fill'], outline=(70, 90, 120), width=3)
    lines = shape['text'].split('\n')
    line_heights = []
    line_widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        line_widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    w = max(line_widths)
    h = sum(line_heights) + (len(lines) - 1) * 6
    tx = shape['xy'][0] + (shape['xy'][2] - shape['xy'][0] - w) / 2
    ty = shape['xy'][1] + (shape['xy'][3] - shape['xy'][1] - h) / 2
    draw.multiline_text((tx, ty), shape['text'], font=subtitle_font, fill=(25, 45, 75), align='center')

# Lines for hierarchy
connects = [((600, 340), (320, 420)), ((600, 340), (660, 420)), ((600, 340), (980, 420)),
            ((320, 520), (320, 620)), ((660, 520), (660, 620)), ((980, 520), (980, 620))]
for p1, p2 in connects:
    draw.line([p1, p2], fill=(50, 80, 120), width=6)
    draw.ellipse([p1[0]-8, p1[1]-8, p1[0]+8, p1[1]+8], fill=(255,255,255), outline=(50,80,120), width=2)
    draw.ellipse([p2[0]-8, p2[1]-8, p2[0]+8, p2[1]+8], fill=(255,255,255), outline=(50,80,120), width=2)

# Additional shapes and layering
draw.rectangle([80, 760, 360, 860], fill=(255, 240, 220), outline=(180, 115, 30), width=4)
draw.ellipse([80, 880, 240, 1040], fill=(220, 245, 255), outline=(10, 90, 140), width=4)
draw.polygon([(420, 820), (620, 820), (520, 950)], fill=(235, 220, 255), outline=(90, 40, 140), width=4)

# Gambar example frame
photo_present = os.path.exists(photo_path)
if photo_present:
    photo = Image.open(photo_path)
    photo.thumbnail((320, 240))
    pw, ph = photo.size
    image.paste(photo, (80 + (320 - pw)//2, 880 + (240 - ph)//2))
    draw.rectangle([80, 880, 400, 1140], outline=(90, 110, 140), width=4)
    draw.text((100, 1148), 'Instruktur & Contoh Gambar', font=subtitle_font, fill=(30, 40, 65))
else:
    draw.text((100, 960), 'Gambar Contoh Tidak Ditemukan', font=subtitle_font, fill=(130, 30, 30))

# Text box visual style
draw.rectangle([460, 820, 1140, 1030], fill=(255, 255, 255), outline=(100, 130, 180), width=4)
text_box = 'Tips: Gunakan Bring to Front / Send to Back untuk mengatur layering.\nGrup objek untuk memindahkan beberapa shape sekaligus.\nSmartArt dapat dibuat dari tab Insert -> SmartArt -> Hierarchy.'
draw.multiline_text((480, 840), text_box, font=subtitle_font, fill=(20, 35, 65))

draw.text((100, 1180), 'Contoh objek: persegi panjang, lingkaran, segitiga, dan teks bergaya WordArt.', font=subtitle_font, fill=(40, 60, 90))

# Save image
os.makedirs(os.path.dirname(img_path), exist_ok=True)
image.save(img_path, 'PNG')

# Create docx with the infographic image
os.makedirs(os.path.dirname(docx_path), exist_ok=True)
doc = Document()
doc.add_heading('Contoh Infografis Word 2019', level=1)
para = doc.add_paragraph('Dokumen ini memuat infografis satu halaman. Contoh elemen: SmartArt hierarki 3 level, WordArt-style title, shapes, gambar, text box visual, dan catatan layering/grouping.')
para.alignment = 0

doc.add_picture(img_path, width=Inches(6.5))
footer = doc.add_paragraph('Lihat gambar untuk representasi desain infografis. Jika Word 2019 tersedia, gunakan fitur Shapes, SmartArt, WordArt, Text Box, Bring to Front, Send to Back, dan Group Objects secara langsung.')
footer.alignment = 0
doc.save(docx_path)
print(f'Created {docx_path}')
